"""
IngestionPipeline — 端到端数据入库流水线。

编排: 文件扫描 → 模态路由 → 预处理 → 嵌入 → 向量库存储。
"""

import logging
import os
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

from tqdm import tqdm

from config.settings import settings
from src.embeddings.jina_embedder import EmbeddingInput, EmbeddingTask, JinaEmbedder, Modality
from src.preprocessing.audio_processor import AudioProcessor
from src.preprocessing.image_processor import ImageProcessor
from src.preprocessing.text_chunker import TextChunker
from src.preprocessing.video_processor import VideoProcessor
from src.storage.schemas import IngestionStats, MediaChunk
from src.storage.vector_store import VectorStore

logger = logging.getLogger(__name__)

# 文件扩展名 → 模态映射
EXTENSION_MAP = {
    # 文本
    ".txt": Modality.TEXT,
    ".md": Modality.TEXT,
    ".markdown": Modality.TEXT,
    ".rst": Modality.TEXT,
    ".py": Modality.TEXT,
    ".js": Modality.TEXT,
    ".ts": Modality.TEXT,
    ".json": Modality.TEXT,
    ".yaml": Modality.TEXT,
    ".yml": Modality.TEXT,
    ".xml": Modality.TEXT,
    ".csv": Modality.TEXT,
    ".html": Modality.TEXT,
    ".tex": Modality.TEXT,
    ".pdf": Modality.TEXT,
    # 文档
    ".docx": Modality.TEXT,
    ".doc": Modality.TEXT,
    ".rtf": Modality.TEXT,
    ".odt": Modality.TEXT,
    ".epub": Modality.TEXT,
    # 图片
    ".jpg": Modality.IMAGE,
    ".jpeg": Modality.IMAGE,
    ".png": Modality.IMAGE,
    ".gif": Modality.IMAGE,
    ".webp": Modality.IMAGE,
    ".bmp": Modality.IMAGE,
    ".tiff": Modality.IMAGE,
    ".tif": Modality.IMAGE,
    # 音频
    ".mp3": Modality.AUDIO,
    ".wav": Modality.AUDIO,
    ".flac": Modality.AUDIO,
    ".ogg": Modality.AUDIO,
    ".m4a": Modality.AUDIO,
    ".opus": Modality.AUDIO,
    ".aac": Modality.AUDIO,
    ".wma": Modality.AUDIO,
    # 视频
    ".mp4": Modality.VIDEO,
    ".mov": Modality.VIDEO,
    ".avi": Modality.VIDEO,
    ".mkv": Modality.VIDEO,
    ".webm": Modality.VIDEO,
    ".flv": Modality.VIDEO,
    ".wmv": Modality.VIDEO,
}


# 媒体文件缓存目录
MEDIA_CACHE_DIR = Path(__file__).resolve().parent.parent.parent / "media_cache"


def _ensure_media_cache() -> Path:
    """确保媒体缓存目录存在。"""
    MEDIA_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return MEDIA_CACHE_DIR


def _save_media(chunk_id: str, data: bytes, suffix: str = ".jpg") -> str:
    """保存媒体文件到缓存，返回相对路径。"""
    cache_dir = _ensure_media_cache()
    fname = f"{chunk_id}{suffix}"
    fpath = cache_dir / fname
    fpath.write_bytes(data)
    return f"media_cache/{fname}"


class IngestionPipeline:
    """端到端数据入库流水线。

    接受文件/目录, 自动检测模态, 经预处理后嵌入并存储到向量库。

    Usage:
        pipeline = IngestionPipeline()
        stats = pipeline.ingest_file("document.pdf")
        stats = pipeline.ingest_directory("./data/")
    """

    def __init__(
        self,
        embedder: Optional[JinaEmbedder] = None,
        vector_store: Optional[VectorStore] = None,
        text_chunker: Optional[TextChunker] = None,
        image_processor: Optional[ImageProcessor] = None,
        audio_processor: Optional[AudioProcessor] = None,
        video_processor: Optional[VideoProcessor] = None,
        batch_size: int = 16,
        incremental: bool = True,
    ):
        self.embedder = embedder or JinaEmbedder()
        self.store = vector_store or VectorStore()
        self.text_chunker = text_chunker or TextChunker(
            chunk_size=settings.text_chunk_size,
            chunk_overlap=settings.text_chunk_overlap,
        )
        self.image_processor = image_processor or ImageProcessor(
            target_size=settings.image_target_size,
            quality=settings.image_quality,
        )
        self.audio_processor = audio_processor or AudioProcessor(
            max_duration_sec=settings.audio_max_duration_sec,
            overlap_sec=settings.audio_overlap_sec,
        )
        self.video_processor = video_processor or VideoProcessor(
            target_size=settings.video_target_size,
            max_frames=settings.video_max_frames,
        )
        self.batch_size = batch_size
        self.incremental = incremental

    # ============================================================
    # 公共 API
    # ============================================================

    def ingest_file(self, file_path: str, display_name: str = "") -> IngestionStats:
        """入库单个文件。

        Args:
            file_path: 文件路径。
            display_name: 显示用的文件名 (优先于实际文件名)。

        Returns:
            IngestionStats 统计信息。
        """
        path = Path(file_path).resolve()
        if not path.exists():
            return IngestionStats(
                files_failed=1,
                errors=[f"文件不存在: {file_path}"],
            )

        src_name = display_name or path.name

        # 增量检查
        if self.incremental and self.store.is_source_indexed(str(path)):
            logger.info(f"文件已索引, 跳过: {src_name}")
            return IngestionStats(files_processed=1)

        t0 = time.perf_counter()
        modality = self._detect_modality(path)

        try:
            if modality == Modality.TEXT:
                chunks = self._ingest_text(path, src_name)
            elif modality == Modality.IMAGE:
                chunks = self._ingest_image(path, src_name)
            elif modality == Modality.AUDIO:
                chunks = self._ingest_audio(path, src_name)
            elif modality == Modality.VIDEO:
                chunks = self._ingest_video(path, src_name)
            else:
                raise ValueError(f"无法检测文件模态: {path}")

            # 嵌入并存储
            if chunks:
                self._embed_and_store(chunks)

            latency = (time.perf_counter() - t0) * 1000

            modality_counts = {}
            for c in chunks:
                m = c.modality.value
                modality_counts[m] = modality_counts.get(m, 0) + 1

            return IngestionStats(
                files_processed=1,
                chunks_created=len(chunks),
                chunks_by_modality=modality_counts,
                total_latency_ms=latency,
            )

        except Exception as e:
            logger.error(f"入库失败 {path}: {e}", exc_info=True)
            return IngestionStats(
                files_failed=1,
                errors=[f"{path.name}: {str(e)}"],
            )

    def ingest_directory(
        self,
        directory: str,
        recursive: bool = True,
    ) -> IngestionStats:
        """批量入库目录中的所有支持文件。

        Args:
            directory: 目录路径。
            recursive: 是否递归处理子目录。

        Returns:
            汇总的 IngestionStats。
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"目录不存在: {directory}")

        # 收集文件
        files = self._collect_files(dir_path, recursive)
        logger.info(f"找到 {len(files)} 个可处理文件")

        # 汇总统计
        total_stats = IngestionStats()

        for fpath in tqdm(files, desc="入库进度"):
            stats = self.ingest_file(str(fpath))
            total_stats.files_processed += stats.files_processed
            total_stats.files_failed += stats.files_failed
            total_stats.chunks_created += stats.chunks_created
            total_stats.total_latency_ms += stats.total_latency_ms
            total_stats.errors.extend(stats.errors)

            for mod, count in stats.chunks_by_modality.items():
                total_stats.chunks_by_modality[mod] = (
                    total_stats.chunks_by_modality.get(mod, 0) + count
                )

        logger.info(
            f"入库完成: {total_stats.files_processed} 个文件, "
            f"{total_stats.chunks_created} 个 chunks, "
            f"{total_stats.files_failed} 个失败"
        )
        return total_stats

    # ============================================================
    # 模态处理
    # ============================================================

    # 文档格式解析器映射
    _DOC_READERS = {
        ".pdf": "_read_pdf",
        ".docx": "_read_docx",
        ".doc": "_read_docx",    # python-docx 也能处理部分 .doc
        ".rtf": "_read_rtf",
        ".odt": "_read_odt",
        ".epub": "_read_epub",
    }

    def _ingest_text(self, path: Path, src_name: str = "") -> List[MediaChunk]:
        """处理文本文件，支持纯文本和多种文档格式。"""
        # 检查文件是否存在且非空
        if not path.exists():
            logger.warning(f"文本文件不存在 (可能已被清理): {path}")
            return []
        file_size = path.stat().st_size
        if file_size == 0:
            logger.warning(f"文本文件为空 (0 字节): {path}")
            return []

        suffix = path.suffix.lower()
        reader = self._DOC_READERS.get(suffix)
        if reader:
            text = getattr(self, reader)(path)
        else:
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                logger.warning(f"文本文件读取失败 ({file_size} 字节): {path}, 错误: {e}")
                return []

        if not text.strip():
            logger.warning(
                f"文本提取为空 ({file_size} 字节, {suffix}): {path.name} — "
                f"{'可能是扫描版 PDF 或图片型文档' if suffix == '.pdf' else '文件可能为二进制内容或无文本层'}"
            )
            return []

        name = src_name or path.name
        chunk_dicts = self.text_chunker.chunk_with_token_limit(
            text=text,
            max_tokens=settings.text_chunk_size,
            source_file=str(path),
        )

        chunks = []
        for cd in chunk_dicts:
            chunks.append(MediaChunk(
                chunk_id=cd["chunk_id"],
                modality=Modality.TEXT,
                content_type="text_chunk",
                source_file=str(path),
                source_file_name=name,
                text_content=cd["text"],
                content_preview=cd["text"][:200],
                original_content_path=str(path),
                embedding_dim=settings.embedding_dimensions,
                extra_metadata=cd["metadata"],
            ))

        return chunks

    def _ingest_image(self, path: Path, src_name: str = "") -> List[MediaChunk]:
        """处理图片文件。"""
        result = self.image_processor.process(str(path))
        chunk_id = result["chunk_id"]

        # 保存媒体文件到缓存
        media_path = _save_media(chunk_id, result["image_bytes"], ".jpg")

        return [MediaChunk(
            chunk_id=chunk_id,
            modality=Modality.IMAGE,
            content_type="image",
            source_file=str(path.resolve()),
            source_file_name=src_name or path.name,
            content_preview=f"[图片: {path.name} ({result['metadata']['original_size'][0]}x{result['metadata']['original_size'][1]})]",
            media_path=media_path,
            thumbnail_base64=result["thumbnail_base64"],
            original_content_path=str(path),
            embedding_dim=settings.embedding_dimensions,
            extra_metadata=result["metadata"],
        )]

    def _ingest_audio(self, path: Path, src_name: str = "") -> List[MediaChunk]:
        """处理音频文件。"""
        segment_dicts = self.audio_processor.process(str(path))

        chunks = []
        for sd in segment_dicts:
            chunk_id = sd["chunk_id"]
            media_path = _save_media(chunk_id, sd["audio_bytes"], ".wav")

            chunks.append(MediaChunk(
                chunk_id=chunk_id,
                modality=Modality.AUDIO,
                content_type="audio_segment",
                source_file=str(path.resolve()),
                source_file_name=src_name or path.name,
                content_preview=sd["text"],
                media_path=media_path,
                start_offset=sd["metadata"]["start_sec"],
                end_offset=sd["metadata"]["end_sec"],
                original_content_path=str(path),
                embedding_dim=settings.embedding_dimensions,
                extra_metadata=sd["metadata"],
            ))

        return chunks

    def _ingest_video(self, path: Path, src_name: str = "") -> List[MediaChunk]:
        """处理视频文件 — 单视频单向量。

        预处理 (缩放+帧率控制) → 整个视频作为一个 VIDEO chunk。
        """
        result = self.video_processor.process(str(path))

        chunk_id = result["chunk_id"]
        media_path = _save_media(chunk_id, result["video_bytes"], ".mp4")

        return [MediaChunk(
            chunk_id=chunk_id,
            modality=Modality.VIDEO,
            content_type="video_segment",
            source_file=str(path.resolve()),
            source_file_name=src_name or path.name,
            content_preview=result["text"],
            media_path=media_path,
            original_content_path=str(path),
            embedding_dim=settings.embedding_dimensions,
            extra_metadata=result["metadata"],
        )]

    # ============================================================
    # 嵌入 & 存储
    # ============================================================

    def _embed_and_store(self, chunks: List[MediaChunk]):
        """批量嵌入并存入向量库。"""
        if not chunks:
            return

        # 构造 EmbeddingInput
        embed_inputs = []
        embeddable_chunks = []  # 与 embed_inputs 一一对应
        for c in chunks:
            if c.modality == Modality.TEXT:
                content = c.text_content or ""
            elif c.media_path:
                media_file = MEDIA_CACHE_DIR / Path(c.media_path).name
                if media_file.exists():
                    content = str(media_file.resolve())
                else:
                    logger.warning(f"媒体文件不存在: {media_file}, 跳过该 chunk")
                    continue
            else:
                logger.warning(f"Chunk {c.chunk_id} 无 media_path 且非文本, 跳过")
                continue

            embed_inputs.append(EmbeddingInput(
                modality=c.modality,
                content=content,
                source_path=c.source_file,
            ))
            embeddable_chunks.append(c)

        if not embed_inputs:
            logger.warning("没有可嵌入的 chunk, 跳过存储")
            return

        # 按模态分组 (Jina API 不允许单次请求混合模态)
        groups: Dict[Modality, List[tuple]] = {}
        for inp, c in zip(embed_inputs, embeddable_chunks):
            groups.setdefault(inp.modality, []).append((inp, c))

        all_embeddings = []
        ordered_chunks = []

        for mod, pairs in groups.items():
            mod_inputs = [p[0] for p in pairs]
            mod_chunks = [p[1] for p in pairs]

            # TEXT 逐条嵌入避免超过 P256 限制; IMAGE/AUDIO/VIDEO 可批量
            effective_batch = 1 if mod == Modality.TEXT else self.batch_size

            for i in tqdm(
                range(0, len(mod_inputs), effective_batch),
                desc=f"嵌入 {mod.value}",
                leave=False,
            ):
                batch = mod_inputs[i:i + effective_batch]
                batch_chunks = mod_chunks[i:i + effective_batch]
                try:
                    result = self.embedder.embed(
                        batch,
                        task=EmbeddingTask.RETRIEVAL_PASSAGE,
                    )
                    all_embeddings.append(result.embeddings)
                    ordered_chunks.extend(batch_chunks)
                except Exception as e:
                    logger.warning(f"嵌入失败 ({mod.value} batch {i}): {e}, 跳过")

        import numpy as np
        if not all_embeddings:
            logger.warning("没有成功嵌入任何 chunk")
            return
        embeddings = np.concatenate(all_embeddings, axis=0)

        # 存入向量库
        self.store.add(embeddings, ordered_chunks)
        logger.debug(f"已存储 {len(ordered_chunks)} 个 chunks")

    # ============================================================
    # 工具方法
    # ============================================================

    @staticmethod
    def _detect_modality(path: Path) -> Modality:
        """检测文件模态。"""
        suffix = path.suffix.lower()
        modality = EXTENSION_MAP.get(suffix)
        if modality:
            return modality
        raise ValueError(f"不支持的文件格式: {suffix}")

    @staticmethod
    def _collect_files(dir_path: Path, recursive: bool) -> List[Path]:
        """收集目录中所有支持的文件。"""
        files = []
        pattern = "**/*" if recursive else "*"
        for fpath in dir_path.glob(pattern):
            if (
                fpath.is_file()
                and fpath.suffix.lower() in EXTENSION_MAP
                and not fpath.name.startswith(".")
            ):
                files.append(fpath)
        return sorted(files)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """读取 PDF 文件文本。

        依次尝试 pdftotext (poppler-utils) 和 PyPDF2。
        对于扫描版/图片型 PDF，两种方式都可能返回空文本。
        """
        import shutil

        # 1. 尝试 pdftotext (本地命令行工具)
        if shutil.which("pdftotext"):
            try:
                import subprocess
                result = subprocess.run(
                    ["pdftotext", str(path), "-"],
                    capture_output=True, text=True, timeout=60,
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
                if result.returncode != 0:
                    logger.debug(f"pdftotext 返回非零: {result.returncode}, stderr: {result.stderr[:200]}")
            except Exception as e:
                logger.debug(f"pdftotext 执行异常: {e}")

        # 2. 回退到 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            num_pages = len(reader.pages)
            logger.info(f"PyPDF2 打开 PDF: {num_pages} 页, 文件: {path.name}")

            texts = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text()
                if t:
                    texts.append(t.strip())

            extracted = "\n\n".join(texts)
            if not extracted.strip():
                logger.warning(
                    f"PDF 文本提取为空 ({num_pages} 页): {path.name} — "
                    "可能是扫描版/图片型 PDF，建议使用 OCR 工具预处理后上传"
                )
            return extracted
        except ImportError:
            logger.warning(
                "PDF 读取需要 pdftotext 或 PyPDF2。"
                "安装: pip install PyPDF2 或 apt install poppler-utils"
            )
            return ""
        except Exception as e:
            logger.error(f"PyPDF2 读取 PDF 失败: {e}")
            return ""

    @staticmethod
    def _read_docx(path: Path) -> str:
        """读取 .docx / .doc 文件文本 (via python-docx)。"""
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也读取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
            if text.strip():
                logger.info(f"python-docx 读取: {path.name}, {len(paragraphs)} 段落/单元格")
            else:
                logger.warning(f"DOCX 文本提取为空: {path.name}")
            return text
        except ImportError:
            logger.warning("读取 .docx 需要 python-docx。安装: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"DOCX 读取失败 ({path.name}): {e}")
            return ""

    @staticmethod
    def _read_rtf(path: Path) -> str:
        """读取 RTF 富文本文件 (via striprtf)。"""
        try:
            from striprtf.striprtf import rtf_to_text
            raw = path.read_text(encoding="utf-8", errors="replace")
            text = rtf_to_text(raw)
            if text.strip():
                logger.info(f"striprtf 读取: {path.name}, {len(text)} 字符")
            else:
                logger.warning(f"RTF 文本提取为空: {path.name}")
            return text
        except ImportError:
            logger.warning("读取 .rtf 需要 striprtf。安装: pip install striprtf")
            return ""
        except Exception as e:
            logger.error(f"RTF 读取失败 ({path.name}): {e}")
            return ""

    @staticmethod
    def _read_odt(path: Path) -> str:
        """读取 .odt OpenDocument 文本文件 (via odfpy)。"""
        try:
            from odf.opendocument import load
            from odf import text as odf_text, teletype
            doc = load(str(path))
            paragraphs = []
            for elem in doc.getElementsByType(odf_text.P):
                t = teletype.extractText(elem)
                if t.strip():
                    paragraphs.append(t.strip())
            text = "\n".join(paragraphs)
            if not text.strip():
                # 回退：作为 ZIP 读取 content.xml
                import zipfile
                with zipfile.ZipFile(str(path)) as z:
                    if "content.xml" in z.namelist():
                        raw = z.read("content.xml").decode("utf-8", errors="replace")
                        import re
                        # 简单提取 <text:p> 中的文本
                        paragraphs = re.findall(r"<text:p[^>]*>(.*?)</text:p>", raw, re.DOTALL)
                        text = "\n".join(re.sub(r"<[^>]+>", "", p).strip() for p in paragraphs)
            if text.strip():
                logger.info(f"ODT 读取: {path.name}, {len(text)} 字符")
            else:
                logger.warning(f"ODT 文本提取为空: {path.name}")
            return text
        except ImportError:
            # odfpy 不可用，尝试 ZIP fallback
            logger.info("odfpy 不可用，使用 ZIP fallback 解析 ODT")
            try:
                import zipfile
                import re
                with zipfile.ZipFile(str(path)) as z:
                    raw = z.read("content.xml").decode("utf-8", errors="replace")
                    paragraphs = re.findall(r"<text:p[^>]*>(.*?)</text:p>", raw, re.DOTALL)
                    text = "\n".join(re.sub(r"<[^>]+>", "", p).strip() for p in paragraphs)
                return text
            except Exception as e2:
                logger.error(f"ODT ZIP fallback 失败 ({path.name}): {e2}")
                return ""
        except Exception as e:
            logger.error(f"ODT 读取失败 ({path.name}): {e}")
            return ""

    @staticmethod
    def _read_epub(path: Path) -> str:
        """读取 .epub 电子书文件 (ZIP + XHTML 解析)。"""
        try:
            import zipfile
            import re
            from html import unescape

            with zipfile.ZipFile(str(path)) as z:
                # 找到所有 XHTML 内容文件
                xhtml_files = [n for n in z.namelist() if n.endswith((".xhtml", ".html", ".htm"))]
                if not xhtml_files:
                    logger.warning(f"EPUB 中未找到 XHTML 内容: {path.name}")
                    return ""

                all_text = []
                for name in sorted(xhtml_files):
                    raw = z.read(name).decode("utf-8", errors="replace")
                    # 提取 <body> 中的文本，去除标签
                    body_match = re.search(r"<body[^>]*>(.*?)</body>", raw, re.DOTALL)
                    content = body_match.group(1) if body_match else raw
                    # 去除 script/style 标签
                    content = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", content, flags=re.DOTALL)
                    # 去除所有 HTML 标签
                    text = re.sub(r"<[^>]+>", " ", content)
                    # 解码 HTML 实体并压缩空白
                    text = unescape(text)
                    text = re.sub(r"\s+", " ", text).strip()
                    if text:
                        all_text.append(text)

                result = "\n\n".join(all_text)
                if result.strip():
                    logger.info(f"EPUB 读取: {path.name}, {len(xhtml_files)} 章节, {len(result)} 字符")
                else:
                    logger.warning(f"EPUB 文本提取为空: {path.name}")
                return result
        except ImportError:
            logger.warning("读取 .epub 需要 Python zipfile 模块 (标准库)")
            return ""
        except Exception as e:
            logger.error(f"EPUB 读取失败 ({path.name}): {e}")
            return ""
